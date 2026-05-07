from pathlib import Path

from docx import Document


def read_docx(path: Path) -> str:
    document = Document(path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"Tabla {table_index}")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)
